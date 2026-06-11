from __future__ import annotations

import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "reports" / "Chapter4.docx"
BACKUP_PATH = ROOT / "reports" / "Chapter4_before_figures.docx"
FIG_ROOT = ROOT / "dissertation_outputs" / "figures"
RESULTS = ROOT / "dissertation_outputs" / "results"
GENERATED = FIG_ROOT / "chapter4_docx_figures"
FINAL_DIR = ROOT / "reports" / "Chapter4_Figures"
GENERATED.mkdir(parents=True, exist_ok=True)
FINAL_DIR.mkdir(parents=True, exist_ok=True)


plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "axes.titlesize": 15,
    "axes.labelsize": 11,
    "xtick.labelsize": 9,
    "ytick.labelsize": 9,
})
sns.set_theme(style="whitegrid")


def save(path: Path) -> Path:
    plt.tight_layout()
    plt.savefig(path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close()
    return path


def flow_diagram(path: Path, title: str, steps: list[str], color: str = "#2E74B5") -> Path:
    fig, ax = plt.subplots(figsize=(13.5, 3.8))
    ax.axis("off")
    ax.set_title(title, weight="bold", pad=18)
    y = 0.52
    xs = [0.08 + i * (0.84 / (len(steps) - 1)) for i in range(len(steps))]
    for i, (x, step) in enumerate(zip(xs, steps)):
        ax.text(
            x,
            y,
            step,
            ha="center",
            va="center",
            color="white",
            fontsize=10,
            bbox=dict(boxstyle="round,pad=0.45,rounding_size=0.08", fc=color, ec="#1F4E79", lw=1.2),
            transform=ax.transAxes,
        )
        if i < len(xs) - 1:
            ax.annotate(
                "",
                xy=(xs[i + 1] - 0.075, y),
                xytext=(x + 0.075, y),
                arrowprops=dict(arrowstyle="->", lw=1.6, color="#555555"),
                xycoords=ax.transAxes,
            )
    return save(path)


def make_results_workflow(path: Path) -> Path:
    fig, ax = plt.subplots(figsize=(15, 8.2))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("#F7F9FC")
    ax.set_facecolor("#F7F9FC")

    ax.text(
        0.5,
        0.955,
        "Chapter 4 Results Workflow",
        ha="center",
        va="center",
        fontsize=21,
        weight="bold",
        color="#1F2937",
    )
    ax.text(
        0.5,
        0.915,
        "From supervised downside-risk prediction to robustness, explainability and cyber-resilience evidence",
        ha="center",
        va="center",
        fontsize=11.5,
        color="#4B5563",
    )

    nodes = [
        ("1", "Data and Target", "OHLCV assets\n15 engineered indicators\n5-day -3% drawdown target", 0.08, 0.66, "#2E74B5"),
        ("2", "Chronological Split", "Train 2010-2019\nValidation 2019-2021\nTest 2021-2023", 0.31, 0.66, "#3A7D44"),
        ("3", "Baseline Models", "Logistic Regression\nRandom Forest\nGradient Boosting fallback", 0.54, 0.66, "#7B4FA1"),
        ("4", "Threshold and Alerts", "Validation-tuned thresholds\nEvent detector\nOR and consensus fusion", 0.77, 0.66, "#C05621"),
        ("5", "Robustness Layer", "Gaussian noise\nFinancial stress scenarios\nAdversarial perturbation", 0.19, 0.28, "#2C7A7B"),
        ("6", "Explainability Layer", "SHAP feature ranking\nPerturbation stability\nFeature audit evidence", 0.42, 0.28, "#9F2F5F"),
        ("7", "Cyber-Resilience Layer", "Cyber-feed simulation\nCase-study mapping\nLayered scorecard", 0.65, 0.28, "#805AD5"),
        ("8", "Evaluation Outputs", "RQ evidence matrix\nDashboards\nDeployment readiness", 0.86, 0.28, "#374151"),
    ]

    def add_node(num: str, title: str, body: str, x: float, y: float, color: str) -> None:
        w, h = 0.17, 0.18
        box = FancyBboxPatch(
            (x - w / 2, y - h / 2),
            w,
            h,
            boxstyle="round,pad=0.018,rounding_size=0.018",
            linewidth=1.4,
            edgecolor=color,
            facecolor="white",
            zorder=2,
        )
        ax.add_patch(box)
        ax.text(
            x - w / 2 + 0.028,
            y + h / 2 - 0.034,
            num,
            ha="center",
            va="center",
            fontsize=10,
            weight="bold",
            color="white",
            bbox=dict(boxstyle="circle,pad=0.23", fc=color, ec=color),
            zorder=3,
        )
        ax.text(x, y + 0.045, title, ha="center", va="center", fontsize=12, weight="bold", color="#111827", zorder=3)
        ax.text(x, y - 0.035, body, ha="center", va="center", fontsize=9.2, color="#374151", linespacing=1.35, zorder=3)

    for node in nodes:
        add_node(*node)

    arrows = [
        ((0.165, 0.66), (0.225, 0.66)),
        ((0.395, 0.66), (0.455, 0.66)),
        ((0.625, 0.66), (0.685, 0.66)),
        ((0.77, 0.57), (0.65, 0.39)),
        ((0.77, 0.57), (0.42, 0.39)),
        ((0.77, 0.57), (0.19, 0.39)),
        ((0.735, 0.28), (0.765, 0.28)),
        ((0.505, 0.28), (0.565, 0.28)),
        ((0.275, 0.28), (0.335, 0.28)),
    ]
    for start, end in arrows:
        ax.add_patch(
            FancyArrowPatch(
                start,
                end,
                arrowstyle="-|>",
                mutation_scale=18,
                linewidth=1.5,
                color="#6B7280",
                shrinkA=2,
                shrinkB=2,
                zorder=1,
            )
        )

    ax.text(
        0.5,
        0.075,
        "Interpretation principle: no single metric is sufficient; evidence is combined across detection, robustness, explanation and operational resilience.",
        ha="center",
        va="center",
        fontsize=10.5,
        color="#1F2937",
        bbox=dict(boxstyle="round,pad=0.45,rounding_size=0.04", fc="#EAF2FB", ec="#B7CBE5"),
    )
    return save(path)


def make_split_distribution(path: Path) -> Path:
    df = pd.read_csv(RESULTS / "time_series_split_summary.csv")
    order = ["Train", "Validation", "Test"]
    df["Split"] = pd.Categorical(df["Split"], categories=order, ordered=True)
    df = df.sort_values("Split")
    fig, ax1 = plt.subplots(figsize=(9, 5.2))
    sns.barplot(data=df, x="Split", y="Rows", color="#4C78A8", ax=ax1)
    ax1.set_title("Dataset Split and Target Distribution", weight="bold")
    ax1.set_ylabel("Rows")
    ax1.set_xlabel("")
    ax2 = ax1.twinx()
    ax2.plot(df["Split"].astype(str), df["Event_Rate"], marker="o", color="#E45756", lw=2.5)
    ax2.set_ylabel("Drawdown event rate")
    ax2.set_ylim(0, max(df["Event_Rate"]) * 1.35)
    for x, y in enumerate(df["Event_Rate"]):
        ax2.text(x, y + 0.006, f"{y:.1%}", ha="center", color="#9B1C1C", fontsize=10)
    return save(path)


def make_event_detector_metrics(path: Path) -> Path:
    df = pd.read_csv(RESULTS / "supervised_event_detector_metrics.csv")
    df = df[df["Split"].eq("Test")].copy()
    df["Model"] = df["Model"].str.replace("Supervised ", "", regex=False).str.replace(" Event Detector", "", regex=False)
    plot = df.melt(id_vars="Model", value_vars=["Precision", "Recall", "F1_Score", "ROC_AUC"], var_name="Metric", value_name="Score")
    fig, ax = plt.subplots(figsize=(10.5, 5.4))
    sns.barplot(data=plot, x="Metric", y="Score", hue="Model", ax=ax)
    ax.set_title("Supervised Event Detector Metrics", weight="bold")
    ax.set_ylim(0, 1.0)
    ax.set_xlabel("")
    ax.set_ylabel("Score")
    ax.legend(title="", loc="upper right")
    return save(path)


def make_shap_stability(path: Path) -> Path:
    df = pd.read_csv(RESULTS / "shap_stability_all_levels_standardized.csv")
    fig, ax = plt.subplots(figsize=(8.8, 5.0))
    ax.plot(df["Perturbation_Level"], df["Spearman_Rho"], marker="o", lw=2.5, color="#2E74B5", label="Spearman rho")
    ax.plot(df["Perturbation_Level"], df["Top_5_Overlap"], marker="s", lw=2.0, color="#E45756", label="Top-5 overlap")
    ax.set_title("SHAP Explanation Stability Across Perturbation Levels", weight="bold")
    ax.set_xlabel("Gaussian perturbation level")
    ax.set_ylabel("Stability score")
    ax.set_ylim(0, 1)
    ax.legend()
    return save(path)


def make_integrated_matrix(path: Path) -> Path:
    rows = [
        ("RQ1", "Baseline metrics", "Detect downside-risk events"),
        ("RQ1", "Threshold tuning", "Recover useful risk ranking"),
        ("RQ2", "Alert fusion", "Balance coverage and alert burden"),
        ("RQ2", "Cyber-feed simulation", "Test suspicious-behaviour coverage"),
        ("RQ3", "Noise and stress tests", "Measure robustness degradation"),
        ("RQ3", "SHAP stability", "Audit explanation reliability"),
        ("RQ3", "Drift monitoring", "Trigger retraining governance"),
    ]
    fig, ax = plt.subplots(figsize=(12, 5.5))
    ax.axis("off")
    table = ax.table(cellText=rows, colLabels=["Research Question", "Evaluation Layer", "Evidence Produced"], loc="center", cellLoc="left")
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 1.55)
    for (r, c), cell in table.get_celld().items():
        if r == 0:
            cell.set_facecolor("#D6E4F0")
            cell.set_text_props(weight="bold")
        elif c == 0:
            cell.set_facecolor("#EDF3F8")
    ax.set_title("Integrated Evaluation Matrix", weight="bold", pad=18)
    return save(path)


def ensure_generated_figures() -> dict[int, Path]:
    return {
        1: make_results_workflow(FINAL_DIR / "Figure_4_1_Results_Workflow.png"),
        2: flow_diagram(
            FINAL_DIR / "Figure_4_2_Supervised_Framework_Pipeline.png",
            "Supervised Framework Pipeline",
            ["OHLCV inputs", "Feature engineering", "Chronological split", "Class-balanced models", "Evaluation", "Analyst-ready outputs"],
            color="#3A7D44",
        ),
        3: make_split_distribution(FINAL_DIR / "Figure_4_3_Dataset_Split_Target_Distribution.png"),
        8: flow_diagram(
            FINAL_DIR / "Figure_4_8_Operational_Alert_Workflow.png",
            "Operational Alert Workflow",
            ["Probability score", "Threshold", "Persistence check", "Alert fusion", "Rationale", "Analyst review"],
            color="#7B4FA1",
        ),
        14: make_event_detector_metrics(FINAL_DIR / "Figure_4_14_Event_Detector_Metrics.png"),
        18: make_shap_stability(FINAL_DIR / "Figure_4_18_SHAP_Stability.png"),
        24: make_integrated_matrix(FINAL_DIR / "Figure_4_24_Integrated_Evaluation_Matrix.png"),
    }


def figure_paths() -> dict[int, Path]:
    generated = ensure_generated_figures()
    existing = {
        4: (FIG_ROOT / "stock_price_trends_multi_asset.png", "Figure_4_4_Multi_Asset_Price_Trends.png"),
        5: (FIG_ROOT / "feature_correlation_heatmap_multi_asset.png", "Figure_4_5_Feature_Correlation_Heatmap.png"),
        6: (FIG_ROOT / "supervised_model_test_metrics.png", "Figure_4_6_Baseline_Model_Test_Metrics.png"),
        7: (FIG_ROOT / "supervised_model_performance_heatmap.png", "Figure_4_7_Supervised_Performance_Heatmap.png"),
        9: (FIG_ROOT / "confusion_matrix_heatmaps_supervised.png", "Figure_4_9_Confusion_Matrix_Heatmaps.png"),
        10: (FIG_ROOT / "precision_recall_curves_supervised.png", "Figure_4_10_Precision_Recall_Curves.png"),
        11: (FIG_ROOT / "roc_curves_supervised.png", "Figure_4_11_ROC_Curves.png"),
        12: (FIG_ROOT / "probability_reliability_diagram.png", "Figure_4_12_Probability_Reliability_Diagram.png"),
        13: (FIG_ROOT / "lift_gains_decile_chart.png", "Figure_4_13_Lift_Gains_Decile_Chart.png"),
        15: (FIG_ROOT / "supervised_f1_drop_under_noise.png", "Figure_4_15_F1_Drop_Under_Gaussian_Noise.png"),
        16: (FIG_ROOT / "supervised_stress_scenario_f1_comparison.png", "Figure_4_16_Stress_Scenario_F1_Comparison.png"),
        17: (FIG_ROOT / "shap_summary_plot_test_sample.png", "Figure_4_17_SHAP_Summary_Plot.png"),
        19: (FIG_ROOT / "cyber_attack_detection_rates_by_type.png", "Figure_4_19_Cyber_Attack_Detection_Rates.png"),
        20: (FIG_ROOT / "cyber_attack_case_study_mapping.png", "Figure_4_20_Cyber_Attack_Case_Study_Mapping.png"),
        21: (FIG_ROOT / "cybersecurity_resilience_radar_chart.png", "Figure_4_21_Cybersecurity_Resilience_Radar.png"),
        22: (FIG_ROOT / "layered_defence_framework_diagram.png", "Figure_4_22_Layered_Defence_Framework.png"),
        23: (FIG_ROOT / "cybersecurity_resilience_summary_dashboard.png", "Figure_4_23_Cybersecurity_Resilience_Dashboard.png"),
        25: (FIG_ROOT / "automated_final_summary_dashboard.png", "Figure_4_25_Automated_Final_Summary_Dashboard.png"),
    }
    copied = {}
    for fig_no, (src, final_name) in existing.items():
        dst = FINAL_DIR / final_name
        shutil.copy2(src, dst)
        copied[fig_no] = dst
    return {**generated, **copied}


def paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_picture_after(paragraph: Paragraph, image_path: Path) -> None:
    picture_para = paragraph_after(paragraph)
    picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_para.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))


def clear_existing_picture_paragraphs(doc: Document) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        has_drawing = bool(paragraph._p.xpath(".//w:drawing"))
        if has_drawing and not paragraph.text.strip():
            paragraph._p.getparent().remove(paragraph._p)
            removed += 1
    return removed


def reset_final_dir() -> None:
    for path in FINAL_DIR.glob("Figure_4_*.png"):
        path.unlink()


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    reset_final_dir()
    paths = figure_paths()
    missing = [str(path) for path in paths.values() if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing figure files:\n" + "\n".join(missing))

    doc = Document(DOCX_PATH)
    removed = clear_existing_picture_paragraphs(doc)
    inserted = 0
    for fig_no in range(25, 0, -1):
        matches = [p for p in doc.paragraphs if p.text.strip().startswith(f"Figure 4.{fig_no}:")]
        if not matches:
            print(f"[WARN] Caption not found for Figure 4.{fig_no}")
            continue
        target_para = matches[-1]
        add_picture_after(target_para, paths[fig_no])
        inserted += 1
    doc.save(DOCX_PATH)
    print(f"Removed {removed} existing picture paragraphs.")
    print(f"Inserted {inserted} figures into {DOCX_PATH}")
    print(f"Backup: {BACKUP_PATH}")
    print(f"Figure folder: {FINAL_DIR}")


if __name__ == "__main__":
    main()
