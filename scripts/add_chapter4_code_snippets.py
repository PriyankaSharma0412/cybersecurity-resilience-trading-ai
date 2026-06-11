from __future__ import annotations

import shutil
from pathlib import Path

import nbformat
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "reports" / "Chapter4_Final_3200_v2.docx"
OUTPUT_DOCX = ROOT / "reports" / "Chapter4_Final_3200_v3_with_more_code_snippets.docx"
OUTPUT_MD = ROOT / "reports" / "Chapter4_Code_Snippets_Appendix.md"
DOWNLOAD_DOCX = Path(r"C:\Users\Priyanka\Downloads") / OUTPUT_DOCX.name
NB_PATH = ROOT / "Updated_Financial_AI_Robustness_Evaluation.ipynb"
PIPELINE_PATH = ROOT / "scripts" / "supervised_pipeline.py"


SNIPPETS = [
    {
        "id": "Code Snippet 4.1",
        "title": "Asset universe and market-data source",
        "cell": 10,
        "ranges": [(1, 23)],
        "script": "scripts/supervised_pipeline.py lines 164-185",
        "insert_after": "Section 4.2 Dataset, Target and Split Summary",
        "why": "Shows the tickers and Yahoo Finance data source used to build the empirical dataset.",
    },
    {
        "id": "Code Snippet 4.2",
        "title": "Engineered feature set",
        "cell": 13,
        "ranges": [(1, 34)],
        "script": "scripts/supervised_pipeline.py lines 185-225",
        "insert_after": "Section 4.2, after feature-engineering discussion",
        "why": "Documents the feature groups used for return, volatility, momentum, drawdown, and technical indicators.",
    },
    {
        "id": "Code Snippet 4.3",
        "title": "Supervised drawdown-risk target creation",
        "cell": 18,
        "ranges": [(2, 6)],
        "script": "scripts/supervised_pipeline.py lines 227-238",
        "insert_after": "Section 4.2, after target-label definition",
        "why": "Defines Target = 1 when the future five-day return is <= -3 percent.",
    },
    {
        "id": "Code Snippet 4.4",
        "title": "Chronological train-validation-test split",
        "cell": 20,
        "ranges": [(2, 20), (25, 37)],
        "script": "scripts/supervised_pipeline.py lines 227-260",
        "insert_after": "Section 4.2, after split-design paragraph",
        "why": "Shows how look-ahead bias is avoided by assigning splits using chronological dates.",
    },
    {
        "id": "Code Snippet 4.5",
        "title": "Model input split into X and y",
        "cell": 22,
        "ranges": [(2, 23)],
        "script": "scripts/supervised_pipeline.py lines 295-302",
        "insert_after": "Section 4.3, before model evaluation",
        "why": "Separates features and target for training, validation, and test evaluation.",
    },
    {
        "id": "Code Snippet 4.6",
        "title": "Gradient Boosting fallback model configuration",
        "cell": 25,
        "ranges": [(1, 40)],
        "script": "scripts/supervised_pipeline.py lines 320-353",
        "insert_after": "Section 4.3, model comparison setup",
        "why": "Explains why XGBoost is used when available and Gradient Boosting fallback otherwise.",
    },
    {
        "id": "Code Snippet 4.7",
        "title": "Classification metrics and confusion-matrix counts",
        "cell": 27,
        "ranges": [(8, 28)],
        "script": "scripts/supervised_pipeline.py lines 360-388",
        "insert_after": "Section 4.3, before Table 4.1",
        "why": "Defines Accuracy, Precision, Recall, F1, ROC-AUC, TP, FP, FN, and TN output rows.",
    },
    {
        "id": "Code Snippet 4.8",
        "title": "Prediction loop and result-table exports",
        "cell": 27,
        "ranges": [(30, 52)],
        "script": "scripts/supervised_pipeline.py lines 365-388",
        "insert_after": "Section 4.3 or 4.4, after baseline results",
        "why": "Shows where predictions, metrics, and confusion matrices are written to CSV.",
    },
    {
        "id": "Code Snippet 4.9",
        "title": "Baseline performance visual selection",
        "cell": 29,
        "ranges": [(1, 25)],
        "script": "scripts/supervised_pipeline.py lines 394-430",
        "insert_after": "Section 4.3, after Figure 4.3",
        "why": "Shows how test-set metrics are selected for baseline figures.",
    },
    {
        "id": "Code Snippet 4.10",
        "title": "Gaussian perturbation robustness testing",
        "cell": 31,
        "ranges": [(2, 39)],
        "script": "scripts/supervised_pipeline.py lines 1558-1578",
        "insert_after": "Section 4.5, robustness paragraph",
        "why": "Adds five noise levels and records prediction flips, F1 drop, AUC drop, and robustness score.",
    },
    {
        "id": "Code Snippet 4.11",
        "title": "Financial stress scenario construction",
        "cell": 33,
        "ranges": [(2, 24), (26, 49)],
        "script": "scripts/supervised_pipeline.py lines 1588-1617",
        "insert_after": "Section 4.5, stress-scenario paragraph",
        "why": "Creates Price Spike, Volume Shock, and Volatility Shock scenarios and compares stressed performance.",
    },
    {
        "id": "Code Snippet 4.12",
        "title": "SHAP feature attribution",
        "cell": 36,
        "ranges": [(2, 27)],
        "script": "scripts/supervised_pipeline.py lines 1777-1790",
        "insert_after": "Section 4.6, SHAP explanation paragraph",
        "why": "Generates the SHAP summary plot and feature-importance table for the strongest tree model.",
    },
    {
        "id": "Code Snippet 4.13",
        "title": "SHAP explanation stability under perturbation",
        "cell": 38,
        "ranges": [(2, 24)],
        "script": "scripts/supervised_pipeline.py lines 1795-1803",
        "insert_after": "Section 4.6, explanation-stability paragraph",
        "why": "Computes Spearman rank correlation and top-five overlap after feature perturbation.",
    },
    {
        "id": "Code Snippet 4.14",
        "title": "Cyber-corrupted data-feed simulation call",
        "cell": 48,
        "ranges": [(1, 26)],
        "script": "scripts/supervised_pipeline.py lines 2637-2688",
        "insert_after": "Section 4.7, cyber-threat simulation paragraph",
        "why": "Runs data-feed poisoning, stale-data injection, gradual drift attack, and coordinated manipulation tests.",
    },
    {
        "id": "Code Snippet 4.15",
        "title": "Attack-surface visual generation",
        "cell": 50,
        "ranges": [(1, 62)],
        "script": "Notebook-only visual cell; related outputs saved under dissertation_outputs/figures/cyber_resilience_visuals",
        "insert_after": "Section 4.7, cyber-resilience figures",
        "why": "Creates visual evidence for attack paths and algorithmic-trading exposure points.",
    },
    {
        "id": "Code Snippet 4.16",
        "title": "Cybersecurity resilience scorecard",
        "cell": 54,
        "ranges": [(1, 43)],
        "script": "scripts/supervised_pipeline.py lines 2699-2742",
        "insert_after": "Section 4.7, scorecard paragraph",
        "why": "Combines robustness, stress, adversarial, SHAP, drift, and cyber detection into resilience layers.",
    },
    {
        "id": "Code Snippet 4.17",
        "title": "Final visual review manifest",
        "cell": 56,
        "ranges": [(1, 32)],
        "script": "scripts/supervised_pipeline.py lines 2513-2541",
        "insert_after": "End of Chapter 4 or figure appendix",
        "why": "Lists the final visual artefacts used for Chapter 4 review.",
    },
    {
        "id": "Code Snippet 4.18",
        "title": "Notebook output manifest",
        "cell": 58,
        "ranges": [(1, 20)],
        "script": "scripts/supervised_pipeline.py lines 2541-2542",
        "insert_after": "End of code appendix",
        "why": "Records generated result and figure outputs for reproducibility.",
    },
]


def shade_paragraph(para, fill="F2F2F2"):
    ppr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def get_notebook_source() -> list:
    return nbformat.read(NB_PATH, as_version=4).cells


def extract_cell_lines(cells, cell_no: int, start: int, end: int) -> str:
    lines = cells[cell_no - 1].get("source", "").splitlines()
    selected = []
    for idx in range(start, min(end, len(lines)) + 1):
        selected.append(f"{idx:03d}: {lines[idx - 1]}")
    return "\n".join(selected)


def add_code_block(doc: Document, code: str):
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        shade_paragraph(p)
    gap = doc.add_paragraph()
    gap.paragraph_format.space_after = Pt(4)


def add_summary_table(doc: Document):
    doc.add_heading("Code Snippet Placement Summary", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Snippet", "Purpose", "Notebook cell and lines", "Pipeline/script reference", "Where to insert", "Reason"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for spec in SNIPPETS:
        cells = table.add_row().cells
        line_parts = [f"Cell {spec['cell']}, lines {a}-{b}" for a, b in spec["ranges"]]
        cells[0].text = spec["id"]
        cells[1].text = spec["title"]
        cells[2].text = "; ".join(line_parts)
        cells[3].text = spec["script"]
        cells[4].text = spec["insert_after"]
        cells[5].text = spec["why"]


def append_snippets(doc: Document):
    cells = get_notebook_source()
    doc.add_page_break()
    doc.add_heading("Appendix A: Code Snippets Used in Chapter 4", level=1)
    doc.add_paragraph(
        "This appendix gives the exact notebook cell and line ranges used as code snippets in Chapter 4. "
        "The snippets support reproducibility and show where the reported target design, model evaluation, robustness testing, explainability, and cyber-resilience outputs come from."
    )
    add_summary_table(doc)
    for spec in SNIPPETS:
        doc.add_heading(f"{spec['id']}: {spec['title']}", level=2)
        doc.add_paragraph(f"Notebook source: Updated_Financial_AI_Robustness_Evaluation.ipynb, Cell {spec['cell']}.")
        doc.add_paragraph(f"Line range considered: {', '.join([f'{a}-{b}' for a, b in spec['ranges']])}.")
        doc.add_paragraph(f"Pipeline/script reference: {spec['script']}.")
        doc.add_paragraph(f"Recommended insertion location: {spec['insert_after']}.")
        doc.add_paragraph(f"Why this snippet is included: {spec['why']}")
        for start, end in spec["ranges"]:
            doc.add_paragraph(f"Cell {spec['cell']}, lines {start}-{end}:")
            add_code_block(doc, extract_cell_lines(cells, spec["cell"], start, end))


def write_markdown_appendix():
    cells = get_notebook_source()
    parts = ["# Appendix A: Code Snippets Used in Chapter 4\n"]
    parts.append("| Snippet | Purpose | Notebook cell and lines | Pipeline/script reference | Where to insert |")
    parts.append("|---|---|---|---|---|")
    for spec in SNIPPETS:
        line_parts = "; ".join([f"Cell {spec['cell']}, lines {a}-{b}" for a, b in spec["ranges"]])
        parts.append(f"| {spec['id']} | {spec['title']} | {line_parts} | {spec['script']} | {spec['insert_after']} |")
    for spec in SNIPPETS:
        parts.append(f"\n## {spec['id']}: {spec['title']}\n")
        parts.append(f"Notebook source: `Updated_Financial_AI_Robustness_Evaluation.ipynb`, Cell {spec['cell']}.")
        parts.append(f"Line range considered: {', '.join([f'{a}-{b}' for a, b in spec['ranges']])}.")
        parts.append(f"Pipeline/script reference: `{spec['script']}`.")
        parts.append(f"Recommended insertion location: {spec['insert_after']}.")
        parts.append(f"Reason: {spec['why']}\n")
        for start, end in spec["ranges"]:
            parts.append(f"```python\n{extract_cell_lines(cells, spec['cell'], start, end)}\n```")
    OUTPUT_MD.write_text("\n".join(parts), encoding="utf-8")


def main():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    shutil.copy2(SOURCE_DOCX, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)
    append_snippets(doc)
    doc.save(OUTPUT_DOCX)
    shutil.copy2(OUTPUT_DOCX, DOWNLOAD_DOCX)
    write_markdown_appendix()
    print(OUTPUT_DOCX)
    print(DOWNLOAD_DOCX)
    print(OUTPUT_MD)
    print(f"Added {len(SNIPPETS)} snippets with exact cell and line ranges.")


if __name__ == "__main__":
    main()
