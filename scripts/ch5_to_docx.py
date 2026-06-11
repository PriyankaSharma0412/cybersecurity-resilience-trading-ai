"""
Convert Chapter5_Final.md to a properly formatted Word document.
Run from project root: .venv\Scripts\python.exe scripts\ch5_to_docx.py
"""
from __future__ import annotations
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH   = Path("reports/Chapter5_Final.md")
DOCX_PATH = Path("reports/Chapter5_Final_v3.docx")
BASE_DIR  = MD_PATH.parent  # images are resolved relative to the .md file


# ---------------------------------------------------------------------------
# Style helpers  (identical to md_to_docx.py)
# ---------------------------------------------------------------------------

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def para_format(para, before=6, after=6, line_spacing=None):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    if line_spacing:
        para.paragraph_format.line_spacing = Pt(line_spacing)


def shade_paragraph(para, fill="F2F2F2"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)


def add_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


def add_inline(para, text: str, base_font="Times New Roman", base_size=12,
               base_bold=False, base_italic=False):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_font(run, base_font, base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            set_font(run, "Courier New", 10)
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            set_font(run, base_font, base_size, italic=True)
        else:
            run = para.add_run(part)
            set_font(run, base_font, base_size, bold=base_bold, italic=base_italic)


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_docx(md_path: Path, docx_path: Path, base_dir: Path):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(3.17)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # Blank
        if not line.strip():
            i += 1
            continue

        # H1
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            h = doc.add_heading(text, level=1)
            h.runs[0].font.name  = "Times New Roman"
            h.runs[0].font.size  = Pt(16)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            para_format(h, before=12, after=6)
            i += 1
            continue

        # H2
        if line.startswith("## "):
            text = line[3:].strip()
            h = doc.add_heading(text, level=2)
            h.runs[0].font.name  = "Times New Roman"
            h.runs[0].font.size  = Pt(13)
            h.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            para_format(h, before=14, after=4)
            i += 1
            continue

        # H3
        if line.startswith("### "):
            text = line[4:].strip()
            h = doc.add_heading(text, level=3)
            if h.runs:
                h.runs[0].font.name  = "Times New Roman"
                h.runs[0].font.size  = Pt(12)
                h.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            para_format(h, before=10, after=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip().startswith("---"):
            i += 1
            continue

        # Code block
        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            for cl in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(cl if cl else " ")
                set_font(run, "Courier New", 9)
                shade_paragraph(p)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                p.paragraph_format.left_indent  = Cm(0.5)
            gap = doc.add_paragraph()
            gap.paragraph_format.space_before = Pt(2)
            gap.paragraph_format.space_after  = Pt(2)
            continue

        # Figure  ![alt](path)
        fig_m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if fig_m:
            alt_text = fig_m.group(1)
            img_rel  = fig_m.group(2)
            img_path = (base_dir / img_rel.replace("/", "\\")).resolve()
            if img_path.exists():
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(img_path), width=Inches(5.8))
                    para_format(p, before=8, after=2)
                except Exception as e:
                    print(f"  [WARN] Could not embed {img_path.name}: {e}")
            else:
                print(f"  [WARN] Image not found: {img_path}")
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(cap, alt_text, base_italic=True, base_size=10)
            para_format(cap, before=0, after=8)
            i += 1
            continue

        # Source line  *Source: ...*
        if re.match(r"^\*Source:", line.strip()):
            src = doc.add_paragraph()
            src.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(src, line.strip().strip("*"), base_italic=True,
                       base_size=9, base_font="Times New Roman")
            para_format(src, before=0, after=6)
            i += 1
            continue

        # Italic-only caption *...*
        if re.match(r"^\*[^*].*[^*]\*$", line.strip()):
            p = doc.add_paragraph()
            add_inline(p, line.strip(), base_italic=True, base_size=10)
            para_format(p, before=0, after=6)
            i += 1
            continue

        # Table
        if line.strip().startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            rows = [r for r in tbl_lines
                    if not re.match(r"^\|[-| :]+\|$", r)]
            if not rows:
                continue
            parsed = [[c.strip() for c in r.strip("|").split("|")] for r in rows]
            n_cols = max(len(r) for r in parsed)
            table  = doc.add_table(rows=len(parsed), cols=n_cols)
            table.style = "Table Grid"
            add_table_borders(table)
            for ri, row_data in enumerate(parsed):
                for ci, cell_text in enumerate(row_data):
                    if ci >= n_cols:
                        break
                    cell = table.rows[ri].cells[ci]
                    cell.text = ""
                    p   = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    set_font(run, "Times New Roman", 10, bold=(ri == 0))
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    if ri == 0:
                        tc   = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd  = OxmlElement("w:shd")
                        shd.set(qn("w:val"),   "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"),  "D6E4F0")
                        tcPr.append(shd)
            doc.add_paragraph()
            continue

        # Table caption  *Table ...*
        if re.match(r"^\*Table", line.strip()):
            p = doc.add_paragraph()
            add_inline(p, line.strip().strip("*"), base_italic=True, base_size=10)
            para_format(p, before=2, after=10)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, line.strip())
        para_format(p, before=0, after=6, line_spacing=14)
        i += 1

    doc.save(docx_path)
    print(f"\nSaved: {docx_path.resolve()}")


if __name__ == "__main__":
    print(f"Converting {MD_PATH} ...")
    build_docx(MD_PATH, DOCX_PATH, BASE_DIR)
    print("Done.")
