from __future__ import annotations

from pathlib import Path

import nbformat
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from create_chapter4_final_2200 import (
    FIG_DIR,
    OUT_DIR,
    RESULT_DIR,
    ROOT,
    add_figure,
    add_table,
    chapter_sections,
    word_count,
)


NB_PATH = ROOT / "Updated_Financial_AI_Robustness_Evaluation.ipynb"
PIPELINE = ROOT / "scripts" / "supervised_pipeline.py"


SNIPPETS = [
    {
        "title": "Code Snippet 4.1: Supervised drawdown-risk target creation",
        "notebook_cell": 18,
        "notebook_lines": (2, 6),
        "script_ref": "scripts/supervised_pipeline.py lines 227-238",
        "placement": "Section 4.2, after the target-label explanation.",
    },
    {
        "title": "Code Snippet 4.2: Chronological train-validation-test split",
        "notebook_cell": 20,
        "notebook_lines": (2, 20),
        "script_ref": "scripts/supervised_pipeline.py lines 227-260",
        "placement": "Section 4.2, after the chronological-split explanation.",
    },
    {
        "title": "Code Snippet 4.3: Classification metrics and confusion-matrix export",
        "notebook_cell": 27,
        "notebook_lines": (8, 28),
        "extra_notebook_lines": (46, 52),
        "script_ref": "scripts/supervised_pipeline.py lines 360-388",
        "placement": "Section 4.3, before Table 4.1.",
    },
    {
        "title": "Code Snippet 4.4: Gaussian perturbation robustness testing",
        "notebook_cell": 31,
        "notebook_lines": (2, 39),
        "script_ref": "scripts/supervised_pipeline.py lines 1558-1578",
        "placement": "Section 4.5, after the first robustness paragraph.",
    },
    {
        "title": "Code Snippet 4.5: Financial stress scenario construction",
        "notebook_cell": 33,
        "notebook_lines": (2, 24),
        "extra_notebook_lines": (32, 49),
        "script_ref": "scripts/supervised_pipeline.py lines 1588-1617",
        "placement": "Section 4.5, after the stress-scenario explanation.",
    },
    {
        "title": "Code Snippet 4.6: SHAP feature importance and stability",
        "notebook_cell": 36,
        "notebook_lines": (11, 27),
        "second_cell": 38,
        "second_cell_lines": (2, 24),
        "script_ref": "scripts/supervised_pipeline.py lines 1777-1803",
        "placement": "Section 4.6, after the SHAP explanation paragraph.",
    },
    {
        "title": "Code Snippet 4.7: Cyber-corrupted data-feed simulation",
        "notebook_cell": 48,
        "notebook_lines": (1, 26),
        "script_ref": "scripts/supervised_pipeline.py lines 2637-2688",
        "placement": "Section 4.7, after the cyber-threat scenario paragraph.",
    },
    {
        "title": "Code Snippet 4.8: Cybersecurity resilience scorecard construction",
        "notebook_cell": 54,
        "notebook_lines": (1, 28),
        "script_ref": "scripts/supervised_pipeline.py lines 2699-2742",
        "placement": "Section 4.7, after the scorecard paragraph.",
    },
]


def get_cell_lines(cell_no: int, start: int, end: int) -> str:
    nb = nbformat.read(NB_PATH, as_version=4)
    cell = nb.cells[cell_no - 1]
    lines = cell.get("source", "").splitlines()
    selected = []
    for idx in range(start, min(end, len(lines)) + 1):
        selected.append(f"{idx:03d}: {lines[idx - 1]}")
    return "\n".join(selected)


def add_code_block(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    para.paragraph_format.space_after = Pt(6)


def add_code_snippets(doc: Document):
    doc.add_page_break()
    doc.add_heading("Code Snippets and Insertion Locations", level=1)
    doc.add_paragraph(
        "These snippets are included for the Chapter 4 evidence trail. The narrative remains 2,200 words; code snippets, tables, and captions are supporting material."
    )
    for spec in SNIPPETS:
        doc.add_heading(spec["title"], level=2)
        doc.add_paragraph(f"Notebook location: Cell {spec['notebook_cell']}, lines {spec['notebook_lines'][0]}-{spec['notebook_lines'][1]}.")
        doc.add_paragraph(f"Pipeline location: {spec['script_ref']}.")
        doc.add_paragraph(f"Suggested insertion point: {spec['placement']}")
        add_code_block(doc, get_cell_lines(spec["notebook_cell"], *spec["notebook_lines"]))
        if "extra_notebook_lines" in spec:
            doc.add_paragraph(f"Continuation from Cell {spec['notebook_cell']}, lines {spec['extra_notebook_lines'][0]}-{spec['extra_notebook_lines'][1]}:")
            add_code_block(doc, get_cell_lines(spec["notebook_cell"], *spec["extra_notebook_lines"]))
        if "second_cell" in spec:
            doc.add_paragraph(f"Continuation from Cell {spec['second_cell']}, lines {spec['second_cell_lines'][0]}-{spec['second_cell_lines'][1]}:")
            add_code_block(doc, get_cell_lines(spec["second_cell"], *spec["second_cell_lines"]))


def add_results_tables(doc: Document):
    doc.add_page_break()
    doc.add_heading("Chapter 4 Result Tables", level=1)

    metrics = pd.read_csv(RESULT_DIR / "supervised_model_metrics.csv")
    test_metrics = metrics[metrics["Split"] == "Test"].rename(columns={"F1_Score": "F1"})
    add_table(
        doc,
        test_metrics[["Model", "Accuracy", "Precision", "Recall", "F1", "ROC_AUC", "TP", "FP", "FN"]],
        ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC_AUC", "TP", "FP", "FN"],
        "Table 4.1: Baseline test-set classifier results.",
    )

    thresholds = pd.read_csv(RESULT_DIR / "validation_refinement_summary.csv")
    thresholds = thresholds.rename(columns={"Optimized_Threshold": "Opt_Threshold", "Test_Optimized_F1": "Opt_Test_F1"})
    add_table(
        doc,
        thresholds[["Model", "Opt_Threshold", "Validation_Best_F1", "Test_Default_F1", "Opt_Test_F1"]],
        ["Model", "Opt_Threshold", "Validation_Best_F1", "Test_Default_F1", "Opt_Test_F1"],
        "Table 4.2: Validation-optimised threshold results.",
    )

    detectors = pd.read_csv(RESULT_DIR / "supervised_event_detector_metrics.csv")
    detectors = detectors[detectors["Split"] == "Test"].rename(columns={"F1_Score": "F1", "Predicted_Events": "Alerts"})
    add_table(
        doc,
        detectors[["Model", "Accuracy", "Precision", "Recall", "F1", "Alerts"]],
        ["Model", "Accuracy", "Precision", "Recall", "F1", "Alerts"],
        "Table 4.3: Supervised event detector results.",
    )

    fusion = pd.read_csv(RESULT_DIR / "unified_alert_fusion_metrics.csv")
    fusion = fusion[fusion["Split"] == "Test"].rename(columns={"F1_Score": "F1", "Predicted_Alerts": "Alerts"})
    keep = [c for c in ["Policy", "Precision", "Recall", "F1", "Alerts"] if c in fusion.columns]
    add_table(doc, fusion[keep], keep, "Table 4.4: Test-set alert fusion results.")

    adversarial = pd.read_csv(RESULT_DIR / "adversarial_attack_results.csv")
    logistic_adv = adversarial[adversarial["Model"] == "Logistic Regression"].rename(columns={"Attacked_ROC_AUC": "Attacked_AUC", "Feature_Budget": "Budget"})
    add_table(
        doc,
        logistic_adv[["Budget", "Baseline_F1", "Attacked_F1", "F1_Drop", "Attacked_AUC", "Targeted_Event_Miss_Rate"]],
        ["Budget", "Baseline_F1", "Attacked_F1", "F1_Drop", "Attacked_AUC", "Targeted_Event_Miss_Rate"],
        "Table 4.5: Logistic Regression adversarial attack results.",
    )

    ttd = pd.read_csv(RESULT_DIR / "time_to_detection_summary.csv")
    base_ttd = ttd[(ttd["Split"] == "Test") & (ttd["Events"] == 343)].rename(columns={"Mean_Time_To_Detection": "Mean_TTD", "Mean_Lead_Time": "Mean_Lead"})
    add_table(
        doc,
        base_ttd[["Model", "Events", "Detected", "Detection_Rate", "Mean_TTD", "Mean_Lead"]],
        ["Model", "Events", "Detected", "Detection_Rate", "Mean_TTD", "Mean_Lead"],
        "Table 4.6: Test-set Time-to-Detection summary.",
    )

    cyber = pd.read_csv(RESULT_DIR / "cyber_attack_data_feed_detection_results.csv")
    cyber_pivot = cyber.pivot(index="Model", columns="Attack_Type", values="Detection_Rate").reset_index()
    add_table(doc, cyber_pivot, list(cyber_pivot.columns), "Table 4.7: Cyber-corrupted data-feed detection rates.")

    scorecard = pd.read_csv(RESULT_DIR / "cybersecurity_resilience_scorecard.csv")
    score_pivot = scorecard.pivot(index="Model", columns="Layer", values="Resilience_Score").reset_index()
    add_table(doc, score_pivot, list(score_pivot.columns), "Table 4.8: Cybersecurity resilience scorecard by layer.")


def add_all_figures(doc: Document):
    doc.add_page_break()
    doc.add_heading("Chapter 4 Figures", level=1)
    figures = [
        (FIG_DIR / "chapter4_visuals" / "01_supervised_framework_pipeline_diagram.png", "Figure 4.1: Supervised framework pipeline."),
        (FIG_DIR / "chapter4_enhanced_visuals" / "chapter4_results_workflow.png", "Figure 4.2: Chapter 4 results workflow."),
        (FIG_DIR / "chapter4_visuals" / "02_dataset_split_and_target_distribution.png", "Figure 4.3: Dataset split and target distribution."),
        (FIG_DIR / "supervised_model_test_metrics.png", "Figure 4.4: Baseline supervised model test metrics."),
        (FIG_DIR / "confusion_matrix_heatmaps_supervised.png", "Figure 4.5: Baseline confusion matrix heatmaps."),
        (FIG_DIR / "chapter4_enhanced_visuals" / "chapter4_operational_alert_workflow.png", "Figure 4.6: Operational alert workflow."),
        (FIG_DIR / "chapter4_visuals" / "04_robustness_and_stress_summary.png", "Figure 4.7: Robustness and stress summary."),
        (FIG_DIR / "supervised_f1_drop_under_noise.png", "Figure 4.8: F1 drop under Gaussian noise."),
        (FIG_DIR / "supervised_stress_scenario_f1_comparison.png", "Figure 4.9: Stress scenario F1 comparison."),
        (FIG_DIR / "shap_summary_plot_test_sample.png", "Figure 4.10: SHAP feature importance summary."),
        (FIG_DIR / "chapter4_visuals" / "05_shap_stability_under_perturbation.png", "Figure 4.11: SHAP stability under perturbation."),
        (FIG_DIR / "cyber_resilience_visuals" / "elaborated_cyber_resilience_dfd.png", "Figure 4.12: Elaborated cyber-resilience DFD."),
        (FIG_DIR / "cyber_resilience_visuals" / "cyber_attack_detection_heatmap.png", "Figure 4.13: Cyber attack detection heatmap."),
        (FIG_DIR / "cyber_resilience_visuals" / "cyber_resilience_layer_heatmap.png", "Figure 4.14: Cybersecurity resilience layer heatmap."),
        (FIG_DIR / "chapter4_enhanced_visuals" / "chapter4_integrated_evaluation_matrix.png", "Figure 4.15: Integrated evaluation matrix."),
    ]
    for path, caption in figures:
        add_figure(doc, path, caption)


def main():
    doc = Document()
    doc.add_heading("Chapter 4: Results, Evaluation and Discussion", level=1)
    wc = sum(word_count(text) for _, text in chapter_sections)
    doc.add_paragraph(f"Narrative word count: {wc} words. Tables, figures, captions, and code snippets are included as supporting material.")

    for heading, text in chapter_sections:
        doc.add_heading(heading, level=2)
        for para in [p.strip() for p in text.strip().split("\n\n") if p.strip()]:
            doc.add_paragraph(para)

    add_results_tables(doc)
    add_all_figures(doc)
    add_code_snippets(doc)

    out = OUT_DIR / "Chapter4_Complete_Updated_All_In_One.docx"
    doc.save(out)
    downloads = Path(r"C:\Users\Priyanka\Downloads") / out.name
    downloads.write_bytes(out.read_bytes())
    print(out)
    print(downloads)
    print("Narrative word count:", wc)


if __name__ == "__main__":
    main()
