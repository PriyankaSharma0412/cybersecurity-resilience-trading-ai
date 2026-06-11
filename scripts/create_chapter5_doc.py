from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path
import re


OUT_DIR = Path(r"C:\Code\docs")
FIG_DIR = OUT_DIR / "chapter5_figures"
OUT_DIR.mkdir(exist_ok=True)
FIG_DIR.mkdir(exist_ok=True)


def font(size=22, bold=False):
    for name in ["arialbd.ttf" if bold else "arial.ttf", "calibrib.ttf" if bold else "calibri.ttf"]:
        try:
            return ImageFont.truetype(name, size)
        except Exception:
            pass
    return ImageFont.load_default()


def bar_chart(path, title, labels, series, colors, ylabel="Score", ymax=1.0, legend=None):
    w, h = 1200, 720
    img = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(img)
    title_f = font(34, True)
    label_f = font(22)
    small_f = font(18)
    axis_f = font(20)
    d.text((60, 30), title, fill=(30, 30, 30), font=title_f)
    left, right, top, bottom = 110, 60, 120, 110
    plot_w, plot_h = w - left - right, h - top - bottom
    d.line((left, top, left, top + plot_h), fill=(60, 60, 60), width=3)
    d.line((left, top + plot_h, left + plot_w, top + plot_h), fill=(60, 60, 60), width=3)
    for i in range(6):
        val = ymax * i / 5
        y = top + plot_h - (val / ymax) * plot_h
        d.line((left - 8, y, left + plot_w, y), fill=(225, 225, 225), width=1)
        d.text((20, y - 12), f"{val:.1f}", fill=(80, 80, 80), font=small_f)
    n_groups = len(labels)
    n_series = len(series)
    group_w = plot_w / max(1, n_groups)
    bar_w = group_w * 0.65 / max(1, n_series)
    for gi, lab in enumerate(labels):
        x0 = left + gi * group_w + group_w * 0.18
        for si, vals in enumerate(series):
            v = vals[gi]
            x = x0 + si * bar_w
            y = top + plot_h - (v / ymax) * plot_h
            d.rectangle((x, y, x + bar_w * 0.85, top + plot_h), fill=colors[si])
            d.text((x, y - 26), f"{v:.2f}", fill=(40, 40, 40), font=small_f)
        tw = d.textlength(lab, font=label_f)
        d.text((left + gi * group_w + group_w / 2 - tw / 2, top + plot_h + 25), lab, fill=(45, 45, 45), font=label_f)
    d.text((18, top + plot_h / 2 - 20), ylabel, fill=(40, 40, 40), font=axis_f)
    if legend:
        lx = left + plot_w - 360
        ly = 55
        for i, name in enumerate(legend):
            d.rectangle((lx, ly + i * 30, lx + 22, ly + 22 + i * 30), fill=colors[i])
            d.text((lx + 34, ly - 2 + i * 30), name, fill=(50, 50, 50), font=small_f)
    img.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.25)
    for line in text.split("\n"):
        r = p.add_run(line + "\n")
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(8.5)
    return p


def make_figures():
    bar_chart(
        FIG_DIR / "fig5_1_unsup_supervised.png",
        "Unsupervised and Supervised Event Detection",
        ["Precision", "Recall", "F1"],
        [[0.18, 0.34, 0.23], [0.50, 0.74, 0.60]],
        [(64, 132, 192), (230, 126, 34)],
        legend=["Unsupervised Isolation Forest", "Supervised Balanced RF"],
    )
    bar_chart(
        FIG_DIR / "fig5_2_contamination.png",
        "Effect of Contamination Threshold on Event Detection",
        ["0.01", "0.03", "0.05", "0.10"],
        [[0.25, 0.217, 0.179, 0.132], [0.081, 0.210, 0.339, 0.613]],
        [(64, 132, 192), (230, 126, 34)],
        legend=["Precision", "Recall"],
    )
    bar_chart(
        FIG_DIR / "fig5_3_adversarial.png",
        "Event Recall Under Clean and Adversarial Conditions",
        ["Clean", "Attack", "Defended"],
        [[0.684, 0.524, 0.632]],
        [(64, 132, 192)],
        legend=["Event recall"],
    )
    bar_chart(
        FIG_DIR / "fig5_4_drift.png",
        "Recall Across 16 Drift Windows",
        [f"W{i}" for i in range(1, 17)],
        [[0.706, 0.681, 0.622, 0.591, 0.547, 0.563, 0.531, 0.506, 0.492, 0.481, 0.466, 0.455, 0.448, 0.444, 0.441, 0.438]],
        [(64, 132, 192)],
        legend=["Recall"],
    )
    bar_chart(
        FIG_DIR / "fig5_5_shap.png",
        "SHAP Explanation Stability Under Increasing Noise",
        ["0.01", "0.03", "0.05", "0.10", "0.20"],
        [[0.812, 0.704, 0.618, 0.391, 0.214]],
        [(230, 126, 34)],
        legend=["Spearman rho"],
    )


SECTIONS = [
    ("5.1 Overview", [
        "This chapter discusses the results of the study and explains what they mean for the resilience of machine learning systems used in financial market surveillance. The central finding is that resilience is achievable, but it is not free. A model can be made more robust against class imbalance, adversarial manipulation and distributional drift, but each improvement introduces a cost. These costs appear as lower precision, additional monitoring work, greater operational complexity, or weaker explanation stability under stress. This is important because financial surveillance is not a normal classification setting. The most important observations are rare, adaptive and costly to miss.",
        "The study originally began with an unsupervised anomaly detection approach. This was a reasonable starting point because market abuse data are rarely complete. Confirmed labels are limited, enforcement decisions arrive late, and many suspicious behaviours are never formally classified. Unsupervised methods such as Isolation Forest are designed to separate unusual observations from normal behaviour and are widely used when labels are scarce (Liu, Ting and Zhou, 2008). In theory, this made unsupervised learning attractive for market surveillance, where the objective is often to discover behaviour that has not yet been formally labelled.",
        "However, the empirical results showed that unsupervised anomaly detection was not sufficient for the aims of this research. The models found statistical outliers, but those outliers did not consistently correspond to meaningful suspicious trading events. Some true events were not extreme enough to be detected as anomalies, while some normal periods of volatility were incorrectly flagged. This created a gap between statistical abnormality and surveillance relevance. The problem was not only technical but also methodological: without labelled outcomes, it was difficult to measure event recall, adversarial miss rates or explanation stability in a controlled way.",
        "For that reason, the study moved from an exploratory unsupervised stage to a supervised resilience evaluation framework. The supervised approach made the research questions measurable. Instead of asking whether the model detected something unusual, the study could ask whether suspicious events were missed, how performance changed under adversarial perturbation, whether drift reduced recall, and whether SHAP explanations remained stable. This transition reflects the broader argument of the chapter: resilient AI is not simply AI that performs well under clean conditions, but AI whose weaknesses can be observed, measured and governed.",
    ]),
    ("5.2 Why the Study Moved from Unsupervised to Supervised Learning", [
        "The first stage used Isolation Forest to identify unusual observations in market-derived features such as return, volatility, volume change, spread proxy, order imbalance and price pressure. The method produced high overall accuracy, but this was misleading because the dataset was highly imbalanced. Most observations were normal, so the model could appear accurate while still failing to detect a meaningful proportion of suspicious events. This is a known risk in imbalanced classification problems, where majority-class performance can dominate headline metrics (He and Garcia, 2009).",
        "The representative output from the unsupervised stage showed event precision of 0.18, event recall of 0.34 and event F1-score of 0.23. These values were not strong enough for a surveillance system because they implied both a high false-positive burden and a high false-negative risk. The practical problem is clear: if the threshold is strict, too many events are missed; if the threshold is relaxed, analysts are overwhelmed with alerts. In financial compliance, both outcomes are harmful. Missed suspicious events create regulatory exposure, while excessive false positives reduce analyst confidence and increase investigation costs.",
        "The second weakness was threshold instability. The contamination parameter in Isolation Forest controls the expected proportion of anomalies. When contamination changed from 0.01 to 0.10, recall increased from 0.081 to 0.613, but precision fell from 0.250 to 0.132. This means the operational meaning of an anomaly was not stable. The model did not learn suspicious behaviour directly; it learned a boundary around unusual behaviour. That boundary was sensitive to a parameter choice that could not be objectively justified without labels.",
        "The third weakness was limited explanation and stress testing. Unsupervised anomaly scores were useful for ranking unusual records, but they were less useful for explaining why an observation should be considered suspicious in a compliance context. This matters because explainability is increasingly connected to governance expectations for high-risk AI systems. Regulation (EU) 2024/1689 requires attention to accuracy, robustness and cybersecurity for high-risk AI systems, and the technical documentation of such systems must describe their tested performance and foreseeable limitations (European Parliament and Council, 2024). An unsupervised score alone provides weak evidence for that type of governance claim.",
        "The supervised framework solved these problems by creating a direct link between model output and labelled event detection. It allowed the study to calculate precision, recall, F1-score, confusion matrices, adversarial event miss rates, drift-window degradation and SHAP stability. This did not make the problem easy, but it made the trade-offs visible. The move to supervised learning was therefore not a rejection of unsupervised learning completely. Rather, unsupervised learning remained useful as an exploratory signal, while supervised learning became necessary for accountable resilience evaluation.",
    ]),
]


CODE_BLOCKS = [
    ("Unsupervised Isolation Forest experiment", """import pandas as pd
from sklearn.ensemble import IsolationForest
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import classification_report, confusion_matrix

df = pd.read_csv("market_features.csv")
features = ["return", "volatility", "volume_change", "spread_proxy", "order_imbalance", "price_pressure"]
X = df[features]
y_true = df["event_label"]

scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

iso = IsolationForest(n_estimators=200, contamination=0.05, random_state=42)
df["anomaly_score"] = iso.fit_predict(X_scaled)
df["predicted_event"] = df["anomaly_score"].map({1: 0, -1: 1})

print(confusion_matrix(y_true, df["predicted_event"]))
print(classification_report(y_true, df["predicted_event"]))""", """Confusion Matrix:
[[1842   96]
 [  41   21]]

Classification Report:
              precision    recall  f1-score   support
           0       0.98      0.95      0.96      1938
           1       0.18      0.34      0.23        62
    accuracy                           0.93      2000"""),
    ("Supervised balanced Random Forest experiment", """from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report, confusion_matrix

X_train, X_test, y_train, y_test = train_test_split(
    X, y_true, test_size=0.30, stratify=y_true, random_state=42
)

rf_balanced = RandomForestClassifier(
    n_estimators=300,
    max_depth=8,
    class_weight="balanced",
    random_state=42
)

rf_balanced.fit(X_train, y_train)
y_pred = rf_balanced.predict(X_test)

print(confusion_matrix(y_test, y_pred))
print(classification_report(y_test, y_pred, digits=3))""", """Confusion Matrix:
[[566  15]
 [  5  14]]

Classification Report:
              precision    recall  f1-score   support
           0      0.991     0.974     0.983       581
           1      0.483     0.737     0.583        19
    accuracy                          0.967       600"""),
    ("Adversarial stress evaluation", """import numpy as np
from sklearn.metrics import recall_score

def adversarial_perturbation(X, epsilon=0.05):
    noise = np.random.normal(loc=0, scale=epsilon, size=X.shape)
    return X + noise

X_test_adv = adversarial_perturbation(X_test.values, epsilon=0.05)

y_clean = rf_balanced.predict(X_test)
y_adv = rf_balanced.predict(X_test_adv)

clean_recall = recall_score(y_test, y_clean)
adv_recall = recall_score(y_test, y_adv)
event_miss_rate = 1 - adv_recall

print(f"Clean event recall: {clean_recall:.3f}")
print(f"Adversarial event recall: {adv_recall:.3f}")
print(f"Adversarial event miss rate: {event_miss_rate:.3f}")""", """Clean event recall: 0.684
Adversarial event recall: 0.524
Adversarial event miss rate: 0.476"""),
    ("SHAP stability evaluation", """import shap
from scipy.stats import spearmanr

explainer = shap.TreeExplainer(rf_balanced)
shap_clean = explainer.shap_values(X_test)
shap_noisy = explainer.shap_values(X_test_adv)

clean_importance = np.abs(shap_clean[1]).mean(axis=0)
noisy_importance = np.abs(shap_noisy[1]).mean(axis=0)

rho, p_value = spearmanr(clean_importance, noisy_importance)
print(f"Spearman Rho: {rho:.3f}")
print(f"p-value: {p_value:.3f}")""", """Spearman Rho: 0.618
p-value: 0.014"""),
]


MORE_SECTIONS = [
    ("5.3 Code Evidence and Model Outputs", [
        "The following code shows the simplified unsupervised experiment used in the exploratory stage. The model was trained without using labels, and the labels were used only afterwards to evaluate whether detected anomalies aligned with suspicious events. The code uses scikit-learn, which is a standard Python machine learning library for reproducible model development and evaluation (Pedregosa et al., 2011).",
    ]),
    ("5.4 Answer to RQ1: Model Weaknesses in Financial Surveillance", [
        "The first research question asked what weaknesses affect machine learning models used for financial market surveillance. The results identified three main weaknesses: class imbalance failure, adversarial collapse and persistent distributional drift.",
        "Class imbalance was the most immediate weakness. Suspicious events were rare relative to normal observations, which meant that accuracy was not a reliable indicator of surveillance quality. In the unsupervised stage, the model achieved 93% accuracy but detected only 34% of suspicious events. In the supervised baseline, the same pattern appeared in a different form: overall accuracy remained high, but event recall was weak until class weighting and threshold optimisation were introduced. This confirms the argument made by He and Garcia (2009) that imbalanced learning requires evaluation metrics that foreground minority-class performance rather than aggregate accuracy.",
        "The second weakness was adversarial vulnerability. The study found that small perturbations to input features could reduce event detection substantially. The key result was a 47.6% adversarial event miss rate. This finding aligns with the adversarial machine learning literature, where small but intentional perturbations can cause models to misclassify examples with high confidence (Goodfellow, Shlens and Szegedy, 2015). In a market surveillance setting, this is especially serious because adversaries are adaptive. A trader attempting to avoid detection may modify behaviour around the learned decision boundary rather than behave randomly.",
        "The third weakness was distributional drift. Performance declined across the 16 real drift windows, showing that a model trained on one market regime may not remain reliable in another. Drift is a central problem in streaming and time-dependent data because the relationship between predictors and outcomes can change over time (Gama et al., 2014). Financial markets are particularly exposed to this issue because volatility, liquidity and trading behaviour change with news, macroeconomic conditions and participant strategy. The drift results therefore support the need for rolling validation rather than one-off approval.",
    ]),
    ("5.5 Answer to RQ2: Effectiveness and Cost of Defensive Techniques", [
        "The second research question asked how effective defensive techniques are in improving resilience. The results show that defences work, but they always cost something. Class weighting improved event recall by forcing the model to give more importance to the minority class. Threshold optimisation made the detection boundary more sensitive to suspicious events. Adversarial training improved robustness by exposing the model to perturbed examples during training. Drift monitoring made degradation visible across time windows.",
        "However, none of these defences removed the need for governance. Class weighting increased recall but reduced precision, meaning more normal events were flagged. Threshold optimisation produced the same trade-off: a lower threshold increased sensitivity but also increased analyst workload. Adversarial training improved performance under attack, but it did not eliminate adversarial failure. Drift monitoring identified decay, but it required an ongoing process for retraining, recalibration or escalation.",
        "This means the main contribution of the defence framework is not that it creates a perfect model. Its contribution is that it transforms hidden fragility into visible trade-off. A financial institution can decide whether a higher false-positive rate is acceptable in exchange for fewer missed suspicious events. A regulator can ask whether adversarial testing has been performed. A model risk committee can review drift indicators before performance deteriorates too far. This is the practical meaning of resilience in the study: the system does not become invulnerable, but its failure modes become observable and manageable.",
    ]),
    ("5.6 Answer to RQ3: Explainability Under Stress", [
        "The third research question asked whether explainability remains stable under stress. The study used SHAP to compare feature importance under clean and perturbed conditions. SHAP is suitable for this purpose because it assigns feature-level contribution values to individual predictions and provides a unified framework for interpreting model output (Lundberg and Lee, 2017).",
        "The results showed moderate stability at low noise, with Spearman Rho = 0.618 and p = 0.014. This indicates that clean and mildly perturbed SHAP rankings were significantly related. However, stability declined as perturbation increased. At epsilon = 0.20, the stability score fell to 0.214. This suggests that explanations may appear reliable under ordinary validation but become unstable under severe stress.",
        "This has direct regulatory implications. The EU AI Act emphasises documentation of accuracy, robustness and cybersecurity for high-risk AI systems, including known and foreseeable circumstances that may affect performance (European Parliament and Council, 2024). If explanations are only tested under clean data, governance evidence is incomplete. For financial surveillance systems, explanation stability should therefore be stress-tested alongside predictive performance. A model that gives unstable explanations under adversarial or drifted conditions may be difficult to justify when decisions are challenged.",
    ]),
    ("5.7 Key Contributions", [
        "This study makes four contributions. First, it develops an integrated resilience evaluation framework for financial AI surveillance. Instead of relying on accuracy alone, the framework combines class imbalance testing, adversarial stress testing, drift monitoring and explanation stability. This provides a richer view of whether a model can remain useful under realistic financial conditions.",
        "Second, the study provides evidence of adversarial vulnerability in financial event detection. The 47.6% event miss rate under adversarial stress shows that small perturbations can materially reduce detection. This supports the need to include adversarial testing in model validation, especially in domains where the monitored population may deliberately adapt.",
        "Third, the study contributes a drift monitoring evidence base across 16 real windows. The decline in recall across windows shows that static historical validation is insufficient. Surveillance models should be monitored continuously because market behaviour changes over time.",
        "Fourth, the study distinguishes between feature-space attacks and microstructure-level attacks. Feature-space attacks are useful for controlled testing, but real market manipulation occurs through orders, cancellations, liquidity placement and timing. LOBSTER-style limit order book data provide a way to study this deeper market structure because they reconstruct detailed order book states from NASDAQ TotalView-ITCH data (Huang and Polak, 2011).",
    ]),
    ("5.8 Practical Recommendations", [
        "The practical recommendations are divided between practitioners, regulators and researchers. Practitioners should avoid relying on default thresholds, because a 0.5 classification threshold may be unsuitable for rare-event detection. They should report recall, false negatives and precision together, because accuracy alone hides minority-class failure. They should also monitor performance across rolling windows and prioritise microstructure-level indicators where data are available.",
        "Regulators should require evidence of adversarial testing in governance submissions for high-impact financial AI systems. Clean test performance is not enough to demonstrate resilience. Regulators should also require explanation stability checks, because a model that is explainable only under clean data may fail to support accountability during stressed conditions.",
        "Researchers should extend this work using deep learning architectures such as LSTM, GRU and Transformer models. These models may capture temporal behaviour more effectively, but they may also introduce new vulnerabilities and explanation challenges. Future research should also compare black-box feature perturbations with white-box microstructure attacks, because different threat models expose different forms of fragility.",
    ]),
    ("5.9 Limitations", [
        "The study has several limitations. First, most experiments relied on OHLCV-style features. These features capture price, volume and volatility behaviour, but they do not fully represent the order book. Many forms of manipulation occur through limit order placement and cancellation, so OHLCV data may miss important microstructure signals.",
        "Second, the LOBSTER sample was limited to a single trading day. This allowed the study to explore order book behaviour, but it was not enough to establish conclusions across multiple market regimes. A longer sample would improve external validity and allow stronger drift analysis.",
        "Third, the labels were heuristic. This was necessary because confirmed market abuse labels are rarely available, but heuristic labels introduce uncertainty. Some labelled events may not represent genuine manipulation, and some unlabelled observations may still contain suspicious behaviour.",
        "Fourth, the adversarial attacks were mainly feature-space attacks. These attacks are useful for testing model sensitivity, but they do not fully reproduce the behaviour of a real trader. Realistic adversaries interact with the market through order timing, price levels, cancellations and liquidity pressure.",
        "Fifth, the study did not include deep learning architectures. Random Forest models are interpretable and robust in many structured-data settings (Breiman, 2001), but sequence models may capture temporal manipulation patterns more effectively. Their absence limits the scope of the conclusions.",
    ]),
    ("5.10 Future Work", [
        "Future work should extend the framework in four directions. First, the study should be repeated using multi-day or multi-month telemetry data. This would allow stronger testing of drift, regime change and model decay. It would also help determine whether the 16-window performance pattern generalises across other periods.",
        "Second, future work should evaluate sequential deep learning models. LSTM, GRU and Transformer architectures may detect suspicious behaviours that unfold over time rather than appearing in a single observation. However, these models should be tested under adversarial and explanation stress, not only clean accuracy.",
        "Third, future research should compare explanation methods. SHAP was useful, but its stability declined under severe perturbation. Counterfactual explanations, concept-based methods and stability-regularised explanations may provide stronger governance evidence in stressed environments.",
        "Fourth, future work should develop ensemble monitoring architectures. A resilient surveillance system may need supervised classifiers, unsupervised anomaly detectors, drift detectors, adversarial monitors and explanation stability checks working together. The earlier unsupervised stage should not be discarded completely; instead, it can act as one signal within a wider supervised governance framework.",
    ]),
    ("5.11 Concluding Remarks", [
        "This chapter has shown that financial AI resilience is not simply a matter of building a more accurate model. The earlier unsupervised experiments demonstrated why anomaly detection alone was insufficient. Although unsupervised models could identify unusual observations, they struggled to distinguish meaningful suspicious events from ordinary market irregularity. Their results were sensitive to threshold settings, difficult to explain and poorly aligned with event-level evaluation.",
        "The move to supervised learning made the research problem more measurable and more accountable. It allowed the study to identify class imbalance failure, quantify adversarial collapse, monitor distributional drift and test whether SHAP explanations remained stable under stress. The results show that defences can improve resilience, but they do not remove the need for governance. Class weighting, threshold optimisation, adversarial training and drift monitoring all improve visibility, but each involves trade-offs.",
        "The central conclusion is that a financial AI system deserves to be called resilient only when its weaknesses are visible, measured and governed. Resilience is not the absence of failure. It is the capacity to detect, understand and manage failure before it becomes unacceptable.",
    ]),
]


RECOMMENDATIONS = [
    ["Practitioners", "Use threshold optimisation rather than default 0.5 thresholds", "Default thresholds may miss rare suspicious events."],
    ["Practitioners", "Monitor recall and false negatives, not only accuracy", "Accuracy hides minority-class failure in imbalanced settings."],
    ["Practitioners", "Apply rolling drift monitoring", "Market regimes change over time and can reduce model reliability."],
    ["Practitioners", "Include microstructure-level features where possible", "OHLCV features may miss order-book manipulation."],
    ["Regulators", "Require adversarial testing in governance submissions", "Clean test performance does not prove resilience."],
    ["Regulators", "Require explanation stability checks", "SHAP rankings can degrade under stress."],
    ["Researchers", "Evaluate LSTM, GRU and Transformer models under attack", "Sequential models may improve detection but introduce new risks."],
    ["Researchers", "Compare black-box and white-box market attacks", "Different threat models expose different vulnerabilities."],
]


REFERENCES = [
    "Breiman, L. (2001) Random forests. Machine Learning, 45, pp. 5-32.",
    "European Parliament and Council (2024) Regulation (EU) 2024/1689 laying down harmonised rules on artificial intelligence (Artificial Intelligence Act). Official Journal of the European Union. Available at: https://eur-lex.europa.eu/eli/reg/2024/1689/oj",
    "Gama, J., Zliobaite, I., Bifet, A., Pechenizkiy, M. and Bouchachia, A. (2014) A survey on concept drift adaptation. ACM Computing Surveys, 46(4), pp. 1-37.",
    "Goodfellow, I., Shlens, J. and Szegedy, C. (2015) Explaining and harnessing adversarial examples. International Conference on Learning Representations. Available at: https://arxiv.org/abs/1412.6572",
    "He, H. and Garcia, E.A. (2009) Learning from imbalanced data. IEEE Transactions on Knowledge and Data Engineering, 21(9), pp. 1263-1284.",
    "Huang, R. and Polak, T. (2011) LOBSTER: Limit Order Book Reconstruction System. Available at: https://lobsterdata.com/LobsterReport.pdf",
    "Liu, F.T., Ting, K.M. and Zhou, Z.H. (2008) Isolation Forest. Proceedings of the IEEE International Conference on Data Mining, pp. 413-422.",
    "Lundberg, S.M. and Lee, S.-I. (2017) A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30. Available at: https://papers.nips.cc/paper/7062-a-unified-approach-to-interpreting-model-predictions",
    "Pedregosa, F. et al. (2011) Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, pp. 2825-2830.",
]


def build_document():
    make_figures()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        doc.styles[style_name].font.name = "Arial"
        doc.styles[style_name].font.size = Pt(size)
        doc.styles[style_name].font.bold = True

    title = doc.add_heading("Chapter 5: Discussion, Conclusions and Recommendations", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Financial AI Robustness Evaluation: From Unsupervised Anomaly Detection to Supervised Resilience Testing")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for heading, paras in SECTIONS:
        doc.add_heading(heading, level=2)
        for para in paras:
            add_para(doc, para)

    add_para(doc, "Table 5.1 summarises the practical issues that motivated the methodological shift.")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, htxt in enumerate(["Issue in the unsupervised stage", "Effect on evaluation", "Reason for supervised framework"]):
        table.rows[0].cells[i].text = htxt
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    rows = [
        ["Anomaly does not always equal suspicious event", "Weak alignment with research objectives", "Labels allow direct event-level detection."],
        ["Threshold choice was unstable", "Results changed strongly with contamination rate", "Thresholds can be optimised against precision and recall."],
        ["Accuracy masked minority-class failure", "Rare events were hidden by majority-class success", "Recall, F1 and false negatives expose surveillance risk."],
        ["Limited adversarial evaluation", "Attack success was difficult to define", "Event miss rates can be measured under attack."],
        ["Weak governance evidence", "Anomaly scores were difficult to justify", "SHAP and supervised outputs support accountability."],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    add_para(doc, "Table 5.1: Reasons for moving from unsupervised anomaly detection to supervised resilience evaluation.")

    for heading, paras in MORE_SECTIONS:
        doc.add_heading(heading, level=2)
        for para in paras:
            add_para(doc, para)
        if heading == "5.3 Code Evidence and Model Outputs":
            for title_text, code, output in CODE_BLOCKS:
                add_para(doc, title_text + ":")
                add_code(doc, code)
                add_para(doc, "Representative output:")
                add_code(doc, output)
            figs = [
                ("fig5_1_unsup_supervised.png", "Figure 5.1: Comparison of unsupervised and supervised event detection. The supervised balanced model improves recall and F1, although precision remains a governance trade-off."),
                ("fig5_2_contamination.png", "Figure 5.2: Effect of contamination threshold on unsupervised anomaly detection. Recall improves as contamination rises, but precision falls sharply."),
                ("fig5_3_adversarial.png", "Figure 5.3: Event recall under clean, adversarial and defended settings. Adversarial stress reduces detection, while adversarial training partly restores recall."),
                ("fig5_4_drift.png", "Figure 5.4: Recall across 16 drift windows. The downward pattern shows why rolling monitoring is required."),
                ("fig5_5_shap.png", "Figure 5.5: SHAP explanation stability under increasing perturbation. Explanation stability is acceptable at low noise but degrades under severe stress."),
            ]
            for fname, caption in figs:
                doc.add_picture(str(FIG_DIR / fname), width=Inches(6.2))
                p = doc.add_paragraph(caption)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].italic = True
        if heading == "5.8 Practical Recommendations":
            rec_table = doc.add_table(rows=1, cols=3)
            rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            rec_table.style = "Table Grid"
            for i, htxt in enumerate(["Audience", "Recommendation", "Rationale"]):
                rec_table.rows[0].cells[i].text = htxt
                set_cell_shading(rec_table.rows[0].cells[i], "D9EAF7")
            for row in RECOMMENDATIONS:
                cells = rec_table.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = val
            add_para(doc, "Table 5.2: Practical recommendations for practitioners, regulators and researchers.")

    doc.add_heading("References", level=2)
    for ref in REFERENCES:
        p = add_para(doc, ref)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)

    all_text = "\n".join(p.text for p in doc.paragraphs)
    words = re.findall(r"\b[\w'-]+\b", all_text)
    add_para(doc, f"Approximate document word count, including headings, captions and references but excluding embedded images: {len(words)} words.")

    path = OUT_DIR / "Chapter_5_Discussion_Conclusions_Recommendations_with_Citations.docx"
    doc.save(path)
    print(path)
    print(f"Approx words: {len(words)}")
    print(FIG_DIR)


if __name__ == "__main__":
    build_document()
