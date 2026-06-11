from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

import nbformat
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from add_chapter4_code_snippets import SNIPPETS


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "reports" / "Chapter4_Final_3200_v2.docx"
OUTPUT_DOCX = ROOT / "reports" / "Chapter4_Final_3200_v4_inline_code_snippets.docx"
DOWNLOAD_DOCX = Path(r"C:\Users\Priyanka\Downloads") / OUTPUT_DOCX.name
NB_PATH = ROOT / "Updated_Financial_AI_Robustness_Evaluation.ipynb"


SECTION_SNIPPETS = {
    "4.2": ["Code Snippet 4.1", "Code Snippet 4.2", "Code Snippet 4.3", "Code Snippet 4.4"],
    "4.3": ["Code Snippet 4.5", "Code Snippet 4.6", "Code Snippet 4.7", "Code Snippet 4.8", "Code Snippet 4.9"],
    "4.5": ["Code Snippet 4.10", "Code Snippet 4.11"],
    "4.6": ["Code Snippet 4.12", "Code Snippet 4.13"],
    "4.7": ["Code Snippet 4.14", "Code Snippet 4.15", "Code Snippet 4.16"],
    "4.8": ["Code Snippet 4.17", "Code Snippet 4.18"],
}


def shade_paragraph(para, fill="F2F2F2"):
    ppr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def extract_cell_lines(cells, cell_no: int, start: int, end: int) -> str:
    lines = cells[cell_no - 1].get("source", "").splitlines()
    return "\n".join(f"{idx:03d}: {lines[idx - 1]}" for idx in range(start, min(end, len(lines)) + 1))


def add_code_block(doc: Document, code: str):
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        shade_paragraph(p)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def build_snippet_elements(doc: Document, cells, snippets: list[dict], section_no: str):
    start_idx = len(doc.element.body)
    doc.add_heading(f"Code evidence for Section {section_no}", level=3)
    intro = doc.add_paragraph(
        "The following code snippets are inserted here because they directly support the results and methods discussed in this section. "
        "Each snippet states the exact notebook cell and line range considered."
    )
    intro.runs[0].font.size = Pt(10)

    for spec in snippets:
        doc.add_heading(f"{spec['id']}: {spec['title']}", level=4)
        meta = doc.add_paragraph()
        meta.add_run("Notebook source: ").bold = True
        meta.add_run(
            f"Updated_Financial_AI_Robustness_Evaluation.ipynb, Cell {spec['cell']}, "
            f"lines {', '.join([f'{a}-{b}' for a, b in spec['ranges']])}. "
        )
        meta.add_run("Pipeline/script reference: ").bold = True
        meta.add_run(spec["script"])
        where = doc.add_paragraph()
        where.add_run("Why included: ").bold = True
        where.add_run(spec["why"])
        for start, end in spec["ranges"]:
            label = doc.add_paragraph(f"Cell {spec['cell']}, lines {start}-{end}:")
            label.runs[0].bold = True
            add_code_block(doc, extract_cell_lines(cells, spec["cell"], start, end))

    end_idx = len(doc.element.body)
    body = doc.element.body
    elements = [body[i] for i in range(start_idx, end_idx)]
    for el in elements:
        body.remove(el)
    return elements


def paragraph_section(paragraph):
    text = paragraph.text.strip()
    if paragraph.style and paragraph.style.name.startswith("Heading 2"):
        for section in ["4.2", "4.3", "4.5", "4.6", "4.7", "4.8"]:
            if text.startswith(section):
                return section
    return None


def find_section_insert_points(doc: Document):
    sections = {}
    current_section = None
    for idx, para in enumerate(doc.paragraphs):
        section = paragraph_section(para)
        if section:
            if current_section and current_section not in sections:
                sections[current_section] = idx - 1
            current_section = section
    if current_section and current_section not in sections:
        sections[current_section] = len(doc.paragraphs) - 1
    return sections


def insert_elements_after_paragraph(paragraph, elements):
    parent = paragraph._p.getparent()
    insert_at = parent.index(paragraph._p)
    for offset, el in enumerate(elements, start=1):
        parent.insert(insert_at + offset, deepcopy(el))


def main():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)

    shutil.copy2(SOURCE_DOCX, OUTPUT_DOCX)
    doc = Document(OUTPUT_DOCX)
    cells = nbformat.read(NB_PATH, as_version=4).cells
    snippets_by_id = {spec["id"]: spec for spec in SNIPPETS}
    insert_points = find_section_insert_points(doc)

    inserted = []
    for section_no, snippet_ids in SECTION_SNIPPETS.items():
        if section_no not in insert_points:
            continue
        snippet_specs = [snippets_by_id[sid] for sid in snippet_ids if sid in snippets_by_id]
        elements = build_snippet_elements(doc, cells, snippet_specs, section_no)
        insert_elements_after_paragraph(doc.paragraphs[insert_points[section_no]], elements)
        inserted.extend(snippet_ids)

    doc.save(OUTPUT_DOCX)
    shutil.copy2(OUTPUT_DOCX, DOWNLOAD_DOCX)
    print(OUTPUT_DOCX)
    print(DOWNLOAD_DOCX)
    print(f"Inserted {len(inserted)} inline snippets inside Chapter 4 sections.")


if __name__ == "__main__":
    main()
