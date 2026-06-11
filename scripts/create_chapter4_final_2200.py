from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "reports"
OUT_DIR.mkdir(exist_ok=True)
RESULT_DIR = ROOT / "dissertation_outputs" / "results"
FIG_DIR = ROOT / "dissertation_outputs" / "figures"


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w'-]+\b", text))


def add_table(doc: Document, df: pd.DataFrame, columns: list[str], title: str):
    doc.add_paragraph(title)
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = col
    for _, row in df[columns].iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            value = row[col]
            if isinstance(value, float):
                cells[i].text = f"{value:.3f}"
            else:
                cells[i].text = str(value)


def add_figure(doc: Document, path: Path, caption: str):
    if path.exists():
        doc.add_picture(str(path), width=Inches(6.25))
    doc.add_paragraph(caption)


chapter_sections = [
("4.1 Introduction and Results Workflow", """
This chapter reports the empirical results from the supervised financial anomaly-detection and cybersecurity-resilience framework developed in Chapter 3. The evaluation is deliberately broader than a normal classifier benchmark. It asks whether models can identify financially significant downside-risk events, whether alerting and anomaly-style monitoring improve detection, and whether model behaviour remains reliable when the input environment becomes noisy, stressed, adversarial, or operationally uncertain. The supervised label defines a positive case when the next five-day return is less than or equal to -3 percent. The models are therefore assessed as early-warning tools for drawdown-risk events, not as trading-profit systems.

The chapter is organised around the pipeline in Figures 4.1 and 4.2. Market data are converted into engineered indicators, split chronologically, used to train Logistic Regression, Random Forest, and Gradient Boosting fallback models, and then evaluated through classification metrics, threshold optimisation, event detection, robustness tests, SHAP explainability, Time-to-Detection, cyber-feed simulation, and deployment-readiness evidence. The main interpretation principle is that no single metric is sufficient. A useful model must catch events, avoid excessive false alarms, remain stable enough under perturbation, and provide evidence that analysts can review.

This structure also prevents the chapter from becoming a list of unrelated outputs. Each result is linked to one of the dissertation questions: predictive performance addresses RQ1, alerting and cyber-monitoring address RQ2, and robustness, explainability, drift, and timing address RQ3. The chapter therefore treats the model as one component inside a wider financial AI control framework.
"""),
("4.2 Dataset, Target and Split Summary", """
The dataset contains daily OHLCV market observations for AAPL, MSFT, NVDA, AMZN, JPM, GS, KO, XOM, SPY, and QQQ. Feature engineering produced return, volatility, momentum, moving-average, volume, skewness, kurtosis, z-score, drawdown, RSI, MACD, MACD signal, and Bollinger Band width variables. These features cover price pressure, trend, risk, distribution shape, and abnormal movement. The target was created using the future five-day return, so the model predicts whether a near-term drawdown event will occur.

Chronological splitting was essential. The training window was used for fitting models, the validation window for threshold tuning, and the test window for final reporting. This avoids look-ahead bias, because future observations are never allowed to influence earlier model development. The test set also contained a higher drawdown-event rate than the training period, making it a more difficult and more realistic evaluation window. The dataset design therefore supports RQ1 directly and also supplies the same prediction outputs used later for robustness, cyber, and operational monitoring analysis.

The target definition is intentionally conservative. A five-day fall of at least 3 percent is large enough to be financially meaningful across liquid equities and ETFs, but frequent enough to provide positive cases for supervised learning. This differs from an unsupervised anomaly label because the event is tied to an observed future outcome rather than only to statistical unusualness. The consequence is that the results should be read as downside-risk detection, not as proof of market manipulation.
"""),
("4.3 Baseline Supervised Model Performance", """
The baseline experiment used the default 0.5 decision threshold. Logistic Regression achieved test accuracy of 0.478, precision of 0.235, recall of 0.745, F1-score of 0.358, and ROC-AUC of 0.612. It detected 764 of the 1,025 true drawdown events, but also produced 2,481 false positives. This made it the strongest baseline early-warning model, because recall is crucial when missed downside events are costly. Its weakness was alert burden.

Random Forest showed the opposite profile. It achieved the highest accuracy, 0.805, and precision, 0.600, but recall was only 0.006. It detected only 6 true events and missed 1,019. This demonstrates why accuracy is misleading for imbalanced financial-risk detection: a model can appear strong by predicting the majority normal class while failing at the monitoring task. Gradient Boosting fallback achieved accuracy of 0.777, precision of 0.231, recall of 0.061, F1-score of 0.097, and ROC-AUC of 0.544. It was less conservative than Random Forest but still missed most events. Overall, RQ1 is answered cautiously: AI models can identify drawdown-risk patterns, but operational usefulness depends on balancing recall, precision, and false-alarm workload.

The confusion counts make the ranking clearer than the headline scores alone. Logistic Regression created the most useful warning coverage because its false negatives were much lower than those of the tree models. Random Forest had only four false positives, but that apparent discipline came at the cost of almost total event non-detection. For a risk-monitoring system, that is not a good trade-off. The preferred baseline model is therefore Logistic Regression, while Random Forest and Gradient Boosting are better treated as comparators or as threshold-tuned alerting components.
"""),
("4.4 Confusion Matrix, Threshold Optimisation and Alert Fusion", """
Confusion matrices clarified the operational meaning of errors. False negatives represent missed drawdown-risk events, while false positives represent alerts without a labelled drawdown. In this context, false positives are not automatically useless because they may still identify stressed market conditions, but they increase analyst workload. False negatives are more dangerous because they allow events to pass without warning. The baseline matrices therefore justify moving beyond default model labels.

Validation-based threshold optimisation changed the model ranking. Logistic Regression used an optimised threshold of 0.594 and achieved an optimised test F1-score of 0.326. Random Forest improved from a default F1-score of 0.012 to 0.334 after its threshold was reduced to 0.077. Gradient Boosting fallback improved from 0.097 to 0.328 using a threshold of 0.092. This shows that the tree models contained useful risk-ranking information even though their default thresholds were too conservative.

Additional supervised event detectors converted probability scores into alerting tools. The Random Forest event detector achieved recall of 0.909 and F1-score of 0.343, but generated 4,411 predicted events. Gradient Boosting fallback achieved recall of 0.767 and F1-score of 0.328, while the Logistic event detector was more balanced. Alert fusion then combined signals. The OR policy maximised recall at 0.957 but produced high alert volume. The consensus policy gave a better F1-score of 0.357 with recall of 0.748. For RQ2, monitoring improves suspicious-event capture, but it needs filtering and human review.

This is why Figure 4.5 is important. It shows that model probabilities must pass through thresholding, persistence checks, alert fusion, rationale generation, and analyst review before becoming operational decisions. The results do not justify automatic trading intervention. They justify a ranked monitoring workflow where severe or repeated alerts are escalated, lower-confidence alerts are reviewed in context, and false-positive burden is measured as a practical cost rather than ignored.
"""),
("4.5 Robustness, Stress, Adversarial and Drift Results", """
Robustness testing examined whether prediction behaviour changed under disturbance. Gaussian noise was added across five perturbation levels. Random Forest had the highest apparent robustness, around 0.998, but this was partly misleading because the baseline model almost never predicted positive events. Logistic Regression was less stable, with average Gaussian robustness around 0.536, but it remained more useful because it detected many events. Gradient Boosting fallback showed high prediction flipping, and its F1 sometimes improved under noise, indicating a poorly aligned default boundary.

Financial stress scenarios tested price spike, volume shock, and volatility shock conditions. Logistic Regression remained comparatively consistent, with worst stressed F1 around 0.346. Random Forest remained stable but weak in detection. Gradient Boosting fallback reacted more strongly, again suggesting threshold sensitivity. Adversarial feature-space testing produced clearer evidence of vulnerability. Under the strongest Logistic Regression attack, F1 dropped from 0.358 to 0.154, attacked ROC-AUC fell to 0.191, and the targeted event miss rate reached 47.6 percent. Drift monitoring also produced frequent retraining triggers across validation and test windows. These results answer RQ3 by showing that model resilience cannot be judged by clean-test metrics alone. A deployable framework requires drift alarms, adversarial review, threshold governance, and rollback or retraining rules.

The key interpretation is that stability and usefulness are different properties. Random Forest appears stable under Gaussian noise because its predictions barely move, but the same model fails to alert on most true events. Logistic Regression is more sensitive, yet it is more valuable as a warning model. Therefore, the desired model is not simply the one with the smallest prediction-flip rate. It is the model whose performance degrades slowly while still detecting a meaningful share of risky cases.
"""),
("4.6 Explainability and Time-to-Detection", """
SHAP was used to explain which features influenced model outputs. The main drivers included MACD signal, volatility, momentum, MACD, drawdown, Bollinger Band width, RSI, skewness, and kurtosis. These are financially meaningful because they describe trend, volatility, abnormal return shape, and drawdown pressure. However, SHAP stability declined as perturbation increased. Spearman rank correlation was moderate at low noise and weaker at stronger disturbance, while top-five feature overlap fell to 0.400 at the highest level. The implication is that SHAP supports analyst review, but explanations should be accompanied by stability and data-quality checks.

Time-to-Detection added an operational measure that static metrics cannot provide. Logistic Regression detected 288 of 343 test drawdown episodes, giving a detection rate of 0.840 and mean lead time of 4.44 trading rows. Random Forest detected only 2 of 343 episodes at the default threshold. Gradient Boosting fallback detected 76. The supervised Random Forest event detector performed much better after threshold tuning, detecting 0.968 of test episodes with mean lead time of 4.64. These results confirm that early-warning value depends not only on whether a model is correct, but also on whether it alerts early enough to support action.

The timing results also explain why threshold-tuned event detectors matter. A static classifier can look weak, but its probability scores may still contain useful early-warning information. When thresholds are tuned on validation data, the detector can raise alerts earlier and more often. This helps the operational workflow, although it again creates the need for alert prioritisation and analyst triage.
"""),
("4.7 Cybersecurity and Operational Resilience", """
The cybersecurity extension tested whether model behaviour changed under cyber-style data disruption. Four data-feed scenarios were simulated: data-feed poisoning, stale-data injection, gradual drift attack, and coordinated manipulation. Detection was uneven. Gradient Boosting fallback responded most strongly to data-feed poisoning, with detection rate of 0.414, while Logistic Regression responded most to gradual drift, with detection rate of 0.319. Coordinated manipulation was weakly detected across all models. These results show that cyber-corrupted inputs can change alert behaviour, but no single model provides full cyber coverage.

Real-world incident mapping and the cyber-resilience scorecard placed these findings into a governance structure. The five layers were Data Integrity, Model Integrity, Drift Monitoring, Explainability, and Operational Response. Random Forest scored strongest on model integrity because of noise stability, while Logistic Regression was more useful as an event-warning model. Drift monitoring was consistently weak, showing frequent distributional change. The scorecard therefore works as a diagnostic tool: it identifies where resilience controls are stronger and where further development is required.

The cyber results must be interpreted carefully. The core dataset is OHLCV market data, not full exchange infrastructure, authentication logs, order-routing records, or network telemetry. For that reason, the cyber experiments are best described as controlled cyber-style data-feed simulations. They strengthen the dissertation by connecting prediction outputs to security-layer thinking, but they do not prove production-grade cyber defence.
"""),
("4.8 Discussion and Research Question Summary", """
RQ1 asked whether AI models can predict financially significant downside-risk events. The answer is yes, but with important limits. Logistic Regression gave the strongest baseline recall and F1-score, while threshold optimisation made Random Forest and Gradient Boosting more competitive. Low precision across active detectors shows that drawdown prediction remains difficult.

RQ2 asked whether anomaly detection and monitoring improve suspicious-behaviour identification. The answer is also qualified. Event detectors, alert fusion, cyber-feed simulation, and telemetry outputs improved coverage, but they increased alert volume and depended partly on proxy evidence rather than full exchange infrastructure data.

RQ3 asked how robust, explainable, and useful the models are under disturbed conditions. The answer is mixed. Clean-test performance did not guarantee robustness. Noise, stress, adversarial, drift, SHAP stability, and Time-to-Detection tests all exposed different weaknesses. The final conclusion is that the framework should be presented as a research-grade decision-support prototype, not an autonomous trading-security system. Its contribution is methodological: it combines prediction, threshold tuning, monitoring, robustness, explainability, cyber simulation, and governance into one reproducible evaluation pipeline.

The strongest overall model choice depends on the operating objective. For baseline early warning, Logistic Regression is preferred because it captures the most drawdown events. For threshold-tuned event detection, Random Forest becomes competitive because its lower threshold produces very high recall. For cyber-feed poisoning, Gradient Boosting fallback is more responsive. The framework contribution is therefore not a single winning model, but a layered evaluation method that shows when each model is useful and where each one fails.

This distinction is important for the dissertation argument. A trading-risk model should not be accepted because one table looks strong, and it should not be rejected because one robustness test looks weak. The evidence has to be read across layers: predictive signal, threshold behaviour, alert timing, explanation stability, drift exposure, cyber-feed sensitivity, and operational response. On that combined basis, the study demonstrates useful but incomplete resilience. The framework closes the research gap by showing how financial AI can be evaluated as a monitored security-sensitive system, while also making clear that live deployment would require richer telemetry, independent validation, and formal governance approval.

This keeps the chapter concise while preserving the empirical logic needed for assessment and clear final examiner review.
"""),
]


code_guide = [
    ["Code Snippet 4.1", "Target creation", "Notebook cell 18, lines 2-6", "scripts/supervised_pipeline.py lines 227-238", "Insert after Section 4.2 first paragraph"],
    ["Code Snippet 4.2", "Chronological split", "Notebook cell 20, lines 2-20", "scripts/supervised_pipeline.py lines 227-260", "Insert after Section 4.2 split explanation"],
    ["Code Snippet 4.3", "Metric and confusion-matrix evaluation", "Notebook cell 27, lines 8-28 and 46-52", "scripts/supervised_pipeline.py lines 360-388", "Insert in Section 4.3 before baseline results"],
    ["Code Snippet 4.4", "Gaussian robustness testing", "Notebook cell 31, lines 2-39", "scripts/supervised_pipeline.py lines 1558-1578", "Insert in Section 4.5"],
    ["Code Snippet 4.5", "Stress scenarios", "Notebook cell 33, lines 2-24 and 32-49", "scripts/supervised_pipeline.py lines 1588-1617", "Insert after Gaussian robustness paragraph"],
    ["Code Snippet 4.6", "SHAP importance and stability", "Notebook cells 36 lines 11-27; 38 lines 2-24", "scripts/supervised_pipeline.py lines 1777-1803", "Insert in Section 4.6"],
    ["Code Snippet 4.7", "Cyber-feed simulation", "Notebook cell 48, lines 1-26", "scripts/supervised_pipeline.py lines 2637-2688", "Insert in Section 4.7"],
    ["Code Snippet 4.8", "Cybersecurity scorecard", "Notebook cell 54, lines 1-28", "scripts/supervised_pipeline.py lines 2699-2742", "Insert after scorecard discussion"],
]


def build_markdown() -> str:
    out = ["# Chapter 4: Results, Evaluation and Discussion\n"]
    out.append("**Word budget used:** " + str(sum(word_count(t) for _, t in chapter_sections)) + " words, excluding tables, captions, and code-insertion guide.\n")
    out.append("**Suggested word split:** 4.1 220; 4.2 230; 4.3 300; 4.4 360; 4.5 330; 4.6 280; 4.7 240; 4.8 240.\n")
    for heading, text in chapter_sections:
        out.append(f"## {heading}\n{text.strip()}\n")
    out.append("## Code Snippet Insertion Guide\n")
    for row in code_guide:
        out.append(f"- **{row[0]} ({row[1]}):** {row[2]}; {row[3]}. Placement: {row[4]}.")
    return "\n".join(out)


def main():
    md = build_markdown()
    md_path = OUT_DIR / "Chapter4_Final_2200.md"
    md_path.write_text(md, encoding="utf-8")

    doc = Document()
    doc.add_heading("Chapter 4: Results, Evaluation and Discussion", level=1)
    doc.add_paragraph(f"Word count note: chapter narrative is {sum(word_count(t) for _, t in chapter_sections)} words, excluding tables, captions, and the separate code-insertion guide.")

    figure_plan = {
        "4.1 Introduction and Results Workflow": [
            (FIG_DIR / "chapter4_visuals" / "01_supervised_framework_pipeline_diagram.png", "Figure 4.1: Supervised framework pipeline."),
            (FIG_DIR / "chapter4_enhanced_visuals" / "chapter4_results_workflow.png", "Figure 4.2: Chapter 4 results workflow."),
        ],
        "4.3 Baseline Supervised Model Performance": [
            (FIG_DIR / "supervised_model_test_metrics.png", "Figure 4.3: Baseline supervised model test metrics."),
        ],
        "4.4 Confusion Matrix, Threshold Optimisation and Alert Fusion": [
            (FIG_DIR / "confusion_matrix_heatmaps_supervised.png", "Figure 4.4: Baseline confusion matrix heatmaps."),
            (FIG_DIR / "chapter4_enhanced_visuals" / "chapter4_operational_alert_workflow.png", "Figure 4.5: Operational alert workflow."),
        ],
        "4.5 Robustness, Stress, Adversarial and Drift Results": [
            (FIG_DIR / "chapter4_visuals" / "04_robustness_and_stress_summary.png", "Figure 4.6: Robustness and stress summary."),
            (FIG_DIR / "supervised_f1_drop_under_noise.png", "Figure 4.7: F1 drop under Gaussian noise."),
        ],
        "4.6 Explainability and Time-to-Detection": [
            (FIG_DIR / "shap_summary_plot_test_sample.png", "Figure 4.8: SHAP feature importance summary."),
            (FIG_DIR / "chapter4_visuals" / "05_shap_stability_under_perturbation.png", "Figure 4.9: SHAP stability under perturbation."),
        ],
        "4.7 Cybersecurity and Operational Resilience": [
            (FIG_DIR / "cyber_resilience_visuals" / "cyber_attack_detection_heatmap.png", "Figure 4.10: Cyber attack detection heatmap."),
            (FIG_DIR / "cyber_resilience_visuals" / "cyber_resilience_layer_heatmap.png", "Figure 4.11: Cybersecurity resilience layer heatmap."),
        ],
        "4.8 Discussion and Research Question Summary": [
            (FIG_DIR / "chapter4_enhanced_visuals" / "chapter4_integrated_evaluation_matrix.png", "Figure 4.12: Integrated evaluation matrix."),
        ],
    }

    for heading, text in chapter_sections:
        doc.add_heading(heading, level=2)
        for para in [p.strip() for p in text.strip().split("\n\n") if p.strip()]:
            doc.add_paragraph(para)
        for path, caption in figure_plan.get(heading, []):
            add_figure(doc, path, caption)

    metrics = pd.read_csv(RESULT_DIR / "supervised_model_metrics.csv")
    test_metrics = metrics[metrics["Split"] == "Test"].copy()
    test_metrics = test_metrics.rename(columns={"F1_Score": "F1"})
    add_table(doc, test_metrics[["Model", "Accuracy", "Precision", "Recall", "F1", "ROC_AUC", "TP", "FP", "FN"]], ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC_AUC", "TP", "FP", "FN"], "Table 4.1: Baseline test-set classifier results.")

    thresholds = pd.read_csv(RESULT_DIR / "validation_refinement_summary.csv")
    thresholds = thresholds.rename(columns={"Optimized_Threshold": "Opt_Threshold", "Test_Optimized_F1": "Opt_Test_F1"})
    add_table(doc, thresholds[["Model", "Opt_Threshold", "Validation_Best_F1", "Test_Default_F1", "Opt_Test_F1"]], ["Model", "Opt_Threshold", "Validation_Best_F1", "Test_Default_F1", "Opt_Test_F1"], "Table 4.2: Validation-optimised threshold results.")

    cyber = pd.read_csv(RESULT_DIR / "cyber_attack_data_feed_detection_results.csv")
    cyber_pivot = cyber.pivot(index="Model", columns="Attack_Type", values="Detection_Rate").reset_index()
    add_table(doc, cyber_pivot, list(cyber_pivot.columns), "Table 4.3: Cyber-corrupted data-feed detection rates.")

    doc.add_page_break()
    doc.add_heading("Code Snippet Insertion Guide", level=1)
    guide_table = doc.add_table(rows=1, cols=5)
    guide_table.style = "Table Grid"
    for i, h in enumerate(["Snippet", "Purpose", "Notebook cell lines", "Script lines", "Placement"]):
        guide_table.rows[0].cells[i].text = h
    for row in code_guide:
        cells = guide_table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

    doc_path = OUT_DIR / "Chapter4_Final_2200_with_Code_Guide.docx"
    doc.save(doc_path)

    guide_doc = Document()
    guide_doc.add_heading("Chapter 4 Code Snippet Insertion Guide", level=1)
    for row in code_guide:
        guide_doc.add_heading(f"{row[0]}: {row[1]}", level=2)
        guide_doc.add_paragraph(f"Notebook location: {row[2]}")
        guide_doc.add_paragraph(f"Pipeline script location: {row[3]}")
        guide_doc.add_paragraph(f"Where to insert: {row[4]}")
    guide_doc.save(OUT_DIR / "Chapter4_Code_Snippet_Insert_Guide.docx")

    print(md_path)
    print(doc_path)
    print(OUT_DIR / "Chapter4_Code_Snippet_Insert_Guide.docx")
    print("Narrative word count:", sum(word_count(t) for _, t in chapter_sections))


if __name__ == "__main__":
    main()
